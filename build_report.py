from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "seoul_growth_projects_report_final.docx"
BLUE = RGBColor(31, 78, 121)
MID_BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(34, 34, 34)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=10.5):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=MID_BLUE if level < 3 else BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_run_font(r, size=10.2, color=DARK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], GRAY_FILL)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(p, 9.5)
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = BLUE
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                set_paragraph_font(p, 9.3)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, side, Inches(1))
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("성장거점형 도심복합개발사업 및 성장잠재권 활성화사업 요약 보고서")
set_run_font(r, size=20, bold=True, color=BLUE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("서울시 자치구 관계자 설명자료 정리 | 2026. 5. 14.")
set_run_font(r, size=11, color=RGBColor(90, 90, 90))

add_heading(doc, "1. 핵심 요약", 1)
add_bullets(
    doc,
    [
        "성장거점형 도심복합개발사업은 민간 참여주체를 확대한 복합개발 모델로, 기존 도심의 고밀복합개발과 신성장거점 조성을 목표로 한다.",
        "성장잠재권 활성화사업은 역세권 밖 주요 간선도로변의 저이용 토지를 대상으로 가로 활성화, 주거환경 개선, 지역 필요시설 공급을 유도한다.",
        "두 사업 모두 용도지역 상향, 용적률 완화, 공공기여 기준을 결합해 사업성을 보완하되 공공성 확보를 전제로 한다.",
        "시범사업은 자치구 추천을 통해 후보지를 발굴하고, 서울시 내부 적정성 검토를 거쳐 최종 대상지를 선정하는 방식으로 추진된다.",
    ],
)

add_heading(doc, "2. 도심복합개발사업: 성장거점형과 주거중심형", 1)
add_heading(doc, "2.1 제도 배경과 목적", 2)
add_body(
    doc,
    "기존 정비사업은 전면수용, 공공 개입, 긴급입법에 따른 세부 규정 미비 등으로 사업 기간 장기화와 사업성 부족 문제가 제기되었다. "
    "새 제도는 민간 주도 사업 참여를 확대하고 일반상업지역까지 용도지역 상향 및 용적률 완화를 허용해 복합개발의 추진 안정성과 실행력을 높이려는 취지다.",
)
add_heading(doc, "2.2 사업 유형 구분", 2)
add_body(
    doc,
    "도심복합개발사업은 설명자료상 성장거점형과 주거중심형으로 유형이 구분된다. 이번 자치구 설명자료와 시범사업 추천 계획은 성장거점형 복합개발에 무게가 실려 있으나, 제도 이해 차원에서는 주거중심형의 위치와 차이를 함께 봐야 한다.",
)
add_table(
    doc,
    ["구분", "성장거점형", "주거중심형"],
    [
        [
            "기능 방향",
            "업무·상업·호텔·지식산업센터·전시공연장 등 중심지 기능과 비주거 복합기능을 강화",
            "주택 공급을 중심으로 하되 상업시설, 생활SOC 등 복합기능을 일부 결합",
        ],
        [
            "입지 요건",
            "5천㎡ 이상, 노후도 미적용. 도심·부도심·생활권 중심지 또는 대중교통 결절지 500m 이내 등 중심성 있는 지역",
            "5천㎡ 이상, 20년 이상 경과 건축물 40% 이상. 면적 과반이 역세권 500m 이내 또는 준공업지역인 곳",
        ],
        [
            "사업 주체",
            "토지등소유자 20명 이내, 신탁, 리츠, 공공시행 등 민간·공공 주체 활용 가능",
            "성장거점형과 동일하게 토지등소유자, 신탁, 리츠, 공공시행 등 활용 가능",
        ],
        [
            "동의 요건",
            "사업시행인가 단계에서 토지등소유자 수 2/3 및 면적 1/2 이상 동의 등 검토",
            "성장거점형과 같은 도심복합개발사업 틀의 동의요건 적용",
        ],
        [
            "규제특례",
            "법적상한 용적률 완화, 건축물 용도제한·건폐율·도시공원녹지 기준 완화 가능",
            "주거지역 1.2배, 준주거지역 1.4배 등 주거공급 중심의 용적률 완화 구조",
        ],
        [
            "공공기여·주택",
            "기반시설·생활인프라·공공시설·국민주택규모 주택 건설·설치, 시·도 조례로 정하는 개발 또는 현금납부 가능",
            "주택 건립 시 용적률 상향 50% 범위에서 공급하되 공공분양 60%, 공공임대 40% 등 주택공급 기준을 중점 검토",
        ],
    ],
    [1.25, 2.65, 2.6],
)
add_heading(doc, "2.3 주요 운영 기준", 2)
add_table(
    doc,
    ["항목", "주요 내용"],
    [
        ["사업 유형", "성장거점형과 주거중심형으로 구분. 성장거점형은 기존 도심 고도화와 신성장거점 조성, 주거중심형은 주택연면적 50% 이상 등 주택공급 중심 구조"],
        ["입지", "도심·부도심·광역중심 또는 환승역 주변지역을 중심으로 검토. 성장거점형은 환승역 반경 500m 등 대중교통 결절성과 중심성을 중시"],
        ["면적", "대상지 규모는 5천㎡ 이상을 기본 기준으로 제시"],
        ["시행 주체", "신탁, 리츠, 토지소유자 등 민간 참여주체 활용. 설명자료의 검토서 예시는 토지소유자 20인 이하 등을 언급"],
        ["동의 요건", "토지등소유자 수 2/3 이상 및 토지면적 1/2 이상 동의 등 사업 추진 동의요건 검토"],
        ["용적률 체계", "기준용적률, 상한용적률, 법적상한용적률의 3단계 체계를 적용. 용도지역 변경 후 조례상 상한과 시행령상 법적상한을 구분"],
        ["공공기여", "용도지역 상향 및 규제특례에 따른 공공기여를 구분해 산정. 평균 표준지공시지가가 낮은 자치구는 증가용적률 기준을 완화하는 방향 제시"],
        ["주거·비주거", "성장거점형 복합개발은 비주거용도 50% 초과 계획 등 중심지 기능 확보 기준을 검토"],
    ],
    [1.5, 5.0],
)
add_heading(doc, "2.4 시범사업 운영", 2)
add_body(
    doc,
    "서울시는 자치구 대상지 추천을 통해 지역 내 고밀복합개발을 선도할 성장거점을 발굴하고, 사업 실현 가능성과 중심지 육성 효과를 함께 고려해 시범사업지를 선정할 계획이다.",
)
add_bullets(
    doc,
    [
        "선정 기준은 정책적 요건, 입지요건, 현황·잠재력, 제외 대상 검토로 구성된다.",
        "정책적 요건에는 관광숙박시설 등 도시성장기반 구축, 지역경제 활성화, 기업유치·전략시설 등이 포함된다.",
        "추진 일정은 2026년 5월 설명회, 6월 시범사업 추천·제안, 7~8월 사업타당성 및 내부 정책회의, 9월 시범사업지 선정 순으로 제시되었다.",
    ],
)

add_heading(doc, "3. 성장잠재권 활성화사업", 1)
add_heading(doc, "3.1 추진 배경", 2)
add_body(
    doc,
    "성장잠재권 활성화사업은 역세권 개발사업 기회가 자치구별로 불균형하게 분포하는 문제를 보완하고, 주요 간선도로변의 저이용 토지와 노후 건축물을 재구조화하기 위해 마련되었다. "
    "설명자료는 폭 35m 이상의 주요 간선도로를 사람들이 활동하고 대중교통 접근성이 높은 기반시설로 보고 이를 ‘성장축’으로 설정한다.",
)
add_heading(doc, "3.2 대상지 기준", 2)
add_table(
    doc,
    ["항목", "운영 기준 요약"],
    [
        ["사업 범위", "폭 35m 이상 간선도로변에 연접한 지역. 용도지역과 타 사업구역 포함 여부에 따라 최종 범위 조정"],
        ["사업 유형", "지구단위계획 또는 도시정비형 재개발 방식 활용"],
        ["면적", "최소 1,500㎡ 이상. 지구단위계획은 5,000㎡ 이하, 도시정비형 재개발은 10,000㎡ 이하"],
        ["용도지역", "제2종일반주거지역(7층 이하 포함), 제3종일반주거지역, 준주거지역 가능"],
        ["접도", "2면 이상이 폭 6m 이상 도로에 접하고, 1면은 폭 35m 이상 간선도로에 대상지 둘레의 1/8 이상 접도"],
        ["제외", "역세권 또는 다른 개발사업 대상지와 중복되거나, 공원·학교·공공주택단지 등 개발이 곤란한 구역은 제외 검토"],
    ],
    [1.5, 5.0],
)
add_heading(doc, "3.3 활성화 전략과 건축 기준", 2)
add_bullets(
    doc,
    [
        "용도지역 변경을 통해 간선도로변 고밀개발을 유도하고, 친환경 건축·관광숙박시설·공개공지 등 용적률 인센티브를 결합한다.",
        "다양한 용도를 도입해 지역 활력을 높이고, 공공기여를 활용해 도로·보육·복지·창업 등 지역 필요시설을 공급한다.",
        "주거계획은 국민주택규모 85㎡ 이하를 원칙으로 하며, 성장잠재권 활성화사업으로 건설하는 주택은 도시형 생활주택 설치를 금지한다.",
        "높이계획은 지구단위계획 기준에 따라 산출된 최고높이의 20% 범위 내 가감 적용이 가능하도록 제시된다.",
        "공개공지, 지속가능형 공동주택, 친환경 건축물, 역사문화 보전 등 공공성과 도시환경 개선 기준을 함께 검토한다.",
    ],
)
add_heading(doc, "3.4 공공기여와 공모 일정", 2)
add_body(
    doc,
    "공공기여율은 증가용적률의 5/10을 기본으로 하며, 자치구 평균 표준지공시지가가 서울시 전체 평균의 60% 이하인 경우 3/10 적용 등 완화 가능성이 제시되었다. "
    "공공시설은 부지 전체가 아니라 증가용적률에 해당하는 용적률을 기준으로 환산하고, 공공주택보다 지역 필요시설과 생활SOC 확보를 우선 검토하는 방향이다.",
)
add_table(
    doc,
    ["일정", "내용"],
    [
        ["2026.5.14.", "성장잠재권 활성화사업 설명 및 추천제안 안내 자치구 교육"],
        ["2026.5.28.", "시범사업 가능 후보지 자치구 추천제안 요청"],
        ["2026.6월", "대상 후보지 접수(자치구 → 서울시)"],
        ["2026.7~8월", "대상 후보지 사업 적정성 검토"],
        ["2026.9월", "성장잠재권 활성화 시범사업 가능 후보지 선정 및 사업 추진"],
    ],
    [1.4, 5.1],
)

add_heading(doc, "4. 두 신규 사업 비교", 1)
add_table(
    doc,
    ["구분", "성장거점형 도심복합개발", "주거중심형 도심복합개발", "성장잠재권 활성화"],
    [
        ["핵심 목표", "대규모 고밀복합개발을 통한 기존 도심 고도화와 신성장거점 조성", "주택공급을 중심으로 한 복합개발과 주거기능 확충", "비역세권 간선도로변의 저이용 토지 활성화와 지역 균형성장"],
        ["공간 초점", "도심·광역중심, 환승역 주변 등 중심지 성격이 강한 지역", "역세권 500m 이내 또는 준공업지역 등 주거공급 잠재지가 있는 지역", "폭 35m 이상 주요 간선도로변, 성장축 주변 지역"],
        ["사업 방식", "도심복합개발법과 서울시 조례·운영기준을 바탕으로 민간 주도 복합개발", "도심복합개발사업 유형 중 주택연면적 50% 이상인 주거 중심 복합개발", "지구단위계획 또는 도시정비형 재개발을 활용한 활성화사업"],
        ["주요 수단", "용도지역 상향, 용적률 완화, 공공기여, 비주거 기능 확보", "용적률 완화, 공공분양·공공임대 등 주택공급 기준, 생활SOC 결합", "용도지역 상향, 용적률 인센티브, 지역 필요시설 공급, 가로환경 개선"],
        ["검토 포인트", "중심지 육성 효과, 사업 실현 가능성, 비주거 기능, 공공기여 적정성", "노후도, 역세권·준공업지역 여부, 주택공급 비율, 공공주택 공급계획", "간선도로 접도, 타 사업 중복 여부, 저이용·노후도, 생활SOC 필요성"],
    ],
    [1.1, 1.8, 1.8, 1.8],
)

add_heading(doc, "5. 기존 역세권 사업들과의 차이점", 1)
add_body(
    doc,
    "서울시의 기존 역세권 사업은 대중교통 결절점 주변을 압축적으로 복합·고밀 개발해 주택공급과 지역필요시설을 확보하는 데 초점을 둔다. "
    "반면 이번 설명자료의 성장거점형 도심복합개발사업과 성장잠재권 활성화사업은 역세권 밖 중심지·간선도로축까지 정책 대상을 확장하고, 민간 참여와 공공기여를 결합해 새로운 사업성 구조를 만들려는 성격이 강하다.",
)
add_heading(doc, "5.1 역세권 활성화사업과 성장잠재권 활성화사업", 2)
add_table(
    doc,
    ["구분", "역세권 활성화사업", "성장잠재권 활성화사업"],
    [
        ["정책 목표", "역세권 복합·고밀개발을 통한 주택공급, 지역필요시설 확충, 저개발·침체지역 활성화", "역세권 밖 주요 간선도로변의 저이용 토지를 활성화하고 비역세권 성장축을 보완"],
        ["공간 범위", "역 주변 대중교통 결절점 중심", "폭 35m 이상 주요 간선도로변과 그 주변 성장잠재권"],
        ["사업 수단", "용도지역 상향, 용적률 증가분 공공기여, 지역필요시설 확보", "용도지역 상향, 용적률 인센티브, 가로환경 개선, 생활SOC 공급"],
        ["공공성", "공공임대주택, 공공임대상가, 생활서비스시설 등 역세권 수요 대응", "도로·보행공간·보육·복지·창업 등 지역 필요시설 확보"],
        ["판단 포인트", "역과의 거리, 역세권 기능, 주택공급 효과", "35m 이상 간선도로 접도, 타 사업 중복 여부, 저이용·노후도, 생활SOC 필요성"],
    ],
    [1.35, 2.55, 2.6],
)
add_heading(doc, "5.2 역세권 활성화사업과 성장거점형 도심복합개발사업", 2)
add_table(
    doc,
    ["구분", "역세권 활성화사업", "성장거점형 도심복합개발사업"],
    [
        ["정책 목표", "역세권 중심의 복합·고밀개발과 지역균형발전", "도심·광역중심 고도화와 신성장거점 조성"],
        ["입지 논리", "역 중심의 대중교통 접근성과 주택·생활서비스 수요", "중심지 기능, 환승역 접근성, 업무·상업·관광 등 전략 기능"],
        ["개발 성격", "주거·상업·생활서비스가 결합된 역세권 복합개발", "비주거 중심기능을 강화하는 대규모 고밀복합개발"],
        ["공공기여", "증가용적률에 따른 공공기여와 지역필요시설 확보", "용도지역 상향 및 규제특례에 따른 공공기여, 중심지 기능 확보"],
        ["판단 포인트", "역세권 기능과 공공기여시설의 지역 필요성", "비주거 기능, 중심지 육성 효과, 사업 실현 가능성"],
    ],
    [1.35, 2.55, 2.6],
)
add_heading(doc, "5.3 장기전세주택과 주거중심형의 직접 비교", 2)
add_body(
    doc,
    "역세권 장기전세주택 사업과 비교할 때의 직접 비교 대상은 성장거점형이 아니라 주거중심형 도심복합개발사업이다. "
    "두 사업 모두 주택공급과 용적률 인센티브를 활용하지만, 장기전세주택은 역세권에서 장기전세주택 공급을 핵심 공공성으로 삼는 반면, 주거중심형은 도심복합개발사업의 한 유형으로 주택연면적 50% 이상, 공공분양·공공임대 공급, 생활SOC 등 복합개발 요소를 함께 검토한다.",
)
add_table(
    doc,
    ["구분", "역세권 장기전세주택", "주거중심형 도심복합개발사업"],
    [
        ["공통점", "역세권 등 입지에서 주택공급과 용적률 인센티브를 결합", "주택공급과 용적률 인센티브를 결합하고 공공주택 공급을 검토"],
        ["핵심 공공성", "장기전세주택 공급이 제도의 핵심 목적", "공공분양·공공임대 등 공공주택과 생활SOC·공공시설을 포함한 복합개발 공공성 확보"],
        ["입지 초점", "역세권 입지를 중심으로 사업성 및 장기전세주택 공급 가능성 검토", "면적 과반이 역세권 500m 이내 또는 준공업지역인 곳 등 주거공급 잠재지 검토"],
        ["개발 성격", "주거공급형 역세권 인센티브 사업 성격", "도심복합개발사업 중 주택연면적 50% 이상인 주거 중심 복합개발 유형"],
        ["검토 포인트", "장기전세주택 공급 규모, 용적률 적용체계, 서울시 매입·공급 구조", "주택연면적 비율, 노후도, 동의요건, 공공분양·공공임대 공급 비율, 생활SOC 계획"],
    ],
    [1.35, 2.55, 2.6],
)
add_heading(doc, "5.4 실무상 구분 기준", 2)
add_bullets(
    doc,
    [
        "후보지가 역세권 내부이고 주택공급·생활서비스시설 확충이 핵심이면 역세권 활성화사업 또는 역세권 장기전세주택과의 정합성을 먼저 검토한다.",
        "후보지가 도심복합개발사업 대상이면서 주택연면적이 50% 이상인 주택공급 중심 계획이라면 성장거점형이 아니라 주거중심형으로 분류될 가능성이 높다.",
        "후보지가 중심지 고도화, 관광숙박·업무·상업 등 비주거 중심기능 확보, 대규모 고밀복합개발이 필요한 곳이면 성장거점형 도심복합개발사업의 적합성이 높다.",
        "후보지가 역세권 밖이지만 폭 35m 이상 주요 간선도로변에 접하고, 저이용 토지·노후 건축물·가로환경 개선 수요가 크면 성장잠재권 활성화사업으로 검토할 여지가 있다.",
        "역세권 사업은 대중교통 결절점 중심의 ‘역 주변 고밀화’ 논리가 강하고, 성장잠재권 활성화사업은 간선도로변의 ‘가로축 재구조화’ 논리가 강하다.",
    ],
)

add_heading(doc, "6. 자치구 검토 체크리스트", 1)
add_bullets(
    doc,
    [
        "후보지가 사업별 입지 기준을 충족하는지 확인: 성장거점형은 중심지·환승역 접근성, 성장잠재권은 35m 이상 간선도로 접도와 타 사업 중복 여부가 핵심이다.",
        "도심복합개발사업 후보지는 먼저 성장거점형인지 주거중심형인지 구분한다. 비주거 중심기능과 중심지 육성 효과가 크면 성장거점형, 주택공급 비중이 높고 주택연면적 50% 이상이면 주거중심형 검토가 필요하다.",
        "역세권 사업으로 처리하는 것이 더 자연스러운 후보지는 아닌지 우선 분류한다. 역세권 내부 후보지는 역세권 활성화사업·역세권 장기전세주택과 신규 사업의 중복 적용 가능성 및 제외 기준을 함께 검토한다.",
        "용도지역 변경 필요성과 변경 단계의 타당성을 검토하고, 조례상 상한과 법적상한 용적률을 구분해 산정한다.",
        "공공기여 시설의 종류, 면적, 비율을 지역 수요와 연결해 제시한다. 도로·보육·복지·창업시설 등 생활SOC 필요성을 근거화하는 것이 중요하다.",
        "주거와 비주거 용도의 비율, 공공주택 공급 유형, 공개공지·보행공간 계획을 함께 검토한다.",
        "대상지 현황사진, 도시관리계획, 기반시설 현황, 종합의견 등 신청 양식에 필요한 자료를 사전에 정리한다.",
    ],
)

add_heading(doc, "7. 종합 시사점", 1)
add_body(
    doc,
    "두 사업은 서울시가 역세권 중심의 정비 프레임을 넘어 중심지와 간선도로축을 함께 활용하려는 정책 방향을 보여준다. "
    "자치구는 단순한 개발 가능지 발굴을 넘어, 해당 지역이 왜 성장거점 또는 성장잠재권으로 기능할 수 있는지, 공공기여가 어떤 지역문제를 해결하는지, 사업 추진 주체와 동의요건을 어떻게 충족할 수 있는지를 함께 제시해야 한다.",
)
add_body(
    doc,
    "특히 시범사업 단계에서는 사업성, 공공성, 속도, 지역 균형성의 균형이 중요하다. 후보지 추천 시에는 입지요건 충족 여부를 우선 확인하고, 공공기여 계획과 가로·보행환경 개선 효과를 구체적으로 제시하는 것이 선정 가능성을 높이는 핵심 대응 방향이다.",
)

add_heading(doc, "8. 참고 기준", 1)
add_bullets(
    doc,
    [
        "서울특별시, 역세권 활성화사업 운영기준(2026.04.30.) 및 사업 안내 페이지",
        "서울특별시, 역세권 장기전세주택 건립 운영기준(2026.3.6. 시행) 및 사업 안내 페이지",
        "서울특별시 도시공간본부, 성장거점형 도심복합개발사업 및 성장잠재권 활성화사업 자치구 관계자 설명자료(2026.5.14.)",
    ],
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("서울시 설명자료 요약 보고서")
set_run_font(r, size=8, color=RGBColor(120, 120, 120))

doc.save(OUT)
print(OUT)
