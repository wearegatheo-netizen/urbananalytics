from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "business_review_structure_template.docx"
BLUE = RGBColor(31, 78, 121)
GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(width)
            margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 12.5, bold=True, color=BLUE)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=10.3)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = header
        shade(c, GRAY)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run(r, size=9.5, bold=True, color=BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run(r, size=9.3)
    set_widths(t, widths)
    doc.add_paragraph()
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, attr, Inches(0.85))

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(72)
title.paragraph_format.space_after = Pt(12)
r = title.add_run("[사업명] 사업검토")
set_run(r, size=22, bold=True, color=BLUE)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("예: 서울 ○○구 ○○동 ○○번지 일대 [사업유형] 사업검토")
set_run(r, size=11, color=RGBColor(90, 90, 90))

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date.add_run("20XX. XX.")
set_run(r, size=11)

doc.add_page_break()

h(doc, "I. 대상지 현황")
body(doc, "대상지의 위치, 면적, 용도지역·지구, 도시관리계획 및 입지 특성을 간결하게 정리한다.")
table(
    doc,
    ["구분", "내용"],
    [
        ["위치", "서울 ○○구 ○○동 ○○번지 일대"],
        ["면적", "○,○○○.○㎡"],
        ["용도·지구", "예: 제○종일반주거지역, 지구단위계획구역, 리모델링지구 등"],
        ["입지 특성", "예: ○○역 반경 ○○m, 폭 ○○m 간선도로 연접, 중심지 위계 등"],
        ["관련 계획", "예: ○○ 지구단위계획, ○○ 활성화구역, 도시관리계획 등"],
    ],
    [1.3, 5.5],
)
table(
    doc,
    ["도면·자료", "첨부 또는 작성 내용"],
    [
        ["토지이용계획도", "대상지 경계, 용도지역·지구, 도시계획시설 표기"],
        ["위치도", "역세권·간선도로·주요 중심지와의 관계 표시"],
        ["현황사진", "전면도로, 주변 건축물, 진출입부, 보행환경 사진"],
    ],
    [1.5, 5.3],
)

h(doc, "II. 사업 요건 검토")
body(doc, "검토 결과는 결론을 먼저 제시하고, 세부 판단은 표로 정리한다.")
bullet(doc, "검토 결론: 대상지는 [사업유형]의 [주요 요건]을 충족/미충족함.")
bullet(doc, "핵심 쟁점: 도로요건, 면적, 노후도, 용도지역, 공공기여, 타 사업 중복 여부 등.")
table(
    doc,
    ["구분", "검토 기준", "검토 결과", "여부"],
    [
        ["도로", "2면 이상 폭 ○m 이상 도로 접도, 최소 1면 폭 ○m 이상", "전면 ○○로 ○m, 측면 ○○로 ○m", "○/×"],
        ["면적", "○,○○○㎡ 이상 ○○,○○○㎡ 이하", "○,○○○.○㎡", "○/×"],
        ["노후도", "20년 이상 경과 건축물 ○/○ 이상 등", "해당/미해당 또는 산정 필요", "○/×/해당없음"],
        ["용도지역", "가능 용도지역 여부", "제○종일반주거지역 등", "○/×"],
        ["복합용도", "업무·관광숙박·비주거 또는 주택 비율 기준", "계획안 기준 ○○% 적용 예정", "○/×"],
        ["공공기여", "증가용적률 기준 공공기여율 및 시설 유형", "공공기여율 ○○%, 시설 ○○ 검토", "검토"],
        ["중복·제외", "타 사업구역 중복 및 개발 불가능 구역 여부", "중복 없음/일부 중복/확인 필요", "○/×"],
    ],
    [1.0, 2.4, 2.4, 0.9],
)

h(doc, "III. 개발검토")
body(doc, "사업 가능성이 있는 경우 밀도, 높이, 용도, 공공기여, 주요 특례를 검토한다.")
table(
    doc,
    ["구분", "계획 내용", "비고"],
    [
        ["용도지역", "기정: ○○지역 / 변경: ○○지역", "위원회 심의 또는 계획 변경 필요"],
        ["건폐율", "○○% 이하", ""],
        ["용적률", "기준(허용) ○○%, 상한 ○○%, 법적상한 ○○%", "인센티브·공공기여 반영"],
        ["가산 항목", "친환경, ZEB, 공개공지, 관광숙박시설 등", "해당 항목만 기재"],
        ["최고높이", "○○m 이하 또는 기준높이 +○○%", "지구단위계획·중심지 위계 검토"],
        ["공공기여", "증가용적률의 ○○% 또는 공공시설 ○○㎡", "시설 종류·위치·면적 구체화"],
        ["용도계획", "주거 ○○%, 비주거 ○○%, 공공시설 ○○%", "사업유형별 기준 충족 여부"],
    ],
    [1.2, 4.1, 1.5],
)
body(doc, "유의사항")
bullet(doc, "상업지역 내 주거복합건물은 주거용 용적률 및 비주거용도 비율 기준을 별도 확인한다.")
bullet(doc, "최고높이는 지구단위계획 수립기준, 블록별 평균높이, 중심지 위계 및 사업별 추가 완화 가능성을 함께 검토한다.")
bullet(doc, "공공기여는 단순 비율뿐 아니라 실제 지역 수요와 설치 가능성까지 확인한다.")

h(doc, "IV. 종합 검토의견")
body(doc, "최종 의견은 적합/조건부 적합/부적합/추가 검토 필요 중 하나로 정리한다.")
table(
    doc,
    ["구분", "검토 의견"],
    [
        ["종합 판단", "예: 도로요건 미충족으로 현 기준상 사업 추진 곤란 / 조건부 검토 가능"],
        ["주요 리스크", "예: 접도요건, 면적, 노후도, 타 사업 중복, 공공기여 부담, 높이 제한"],
        ["보완 필요사항", "예: 도로 확보, 구역계 조정, 용도계획 보완, 공공기여 시설 구체화"],
        ["후속 검토", "예: 관계부서 협의, 지구단위계획 정합성 검토, 사업성 분석, 법령 검토"],
    ],
    [1.4, 5.4],
)

doc.save(OUT)
print(OUT)
